"""
🟣 Rally Agent — RAG Pipeline
Retrieval-Augmented Generation with document ingestion, chunking,
vector search, hybrid retrieval, and source citation.

Supports: PDF, DOCX, Markdown, HTML, plain text, code files.
Uses sentence-transformers for embeddings with TF-IDF fallback.
ChromaDB for persistent vector storage, or in-memory store.
"""

from __future__ import annotations

import hashlib
import json
import logging
import math
import os
import re
import struct
import time
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import (
    Any,
    Dict,
    List,
    Literal,
    Optional,
    Protocol,
    Sequence,
    Tuple,
    TypeAlias,
    Union,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional dependencies
# ---------------------------------------------------------------------------

try:
    import numpy as np

    HAS_NUMPY = True
except ImportError:
    np = None  # type: ignore[assignment]
    HAS_NUMPY = False

try:
    from sentence_transformers import SentenceTransformer

    HAS_SENTENCE_TRANSFORMERS = True
except ImportError:
    SentenceTransformer = None  # type: ignore[assignment,misc]
    HAS_SENTENCE_TRANSFORMERS = False

try:
    import chromadb
    from chromadb.config import Settings as ChromaSettings

    HAS_CHROMADB = True
except ImportError:
    chromadb = None  # type: ignore[assignment]
    HAS_CHROMADB = False

# ---------------------------------------------------------------------------
# Type aliases
# ---------------------------------------------------------------------------

EmbeddingVector: TypeAlias = List[float]
JsonDict: TypeAlias = Dict[str, Any]

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DEFAULT_EMBEDDING_MODEL = "all-MiniLM-L6-v2"
DEFAULT_CHUNK_SIZE = 512
DEFAULT_CHUNK_OVERLAP = 64
DEFAULT_TOP_K = 5
EMBEDDING_INDEX_EXT = ".embidx"

# Supported file types and their MIME categories
SUPPORTED_EXTENSIONS: Dict[str, str] = {
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".html": "html",
    ".htm": "html",
    ".pdf": "pdf",
    ".docx": "docx",
    ".py": "code",
    ".js": "code",
    ".ts": "code",
    ".tsx": "code",
    ".jsx": "code",
    ".go": "code",
    ".rs": "code",
    ".rb": "code",
    ".java": "code",
    ".c": "code",
    ".cpp": "code",
    ".h": "code",
    ".hpp": "code",
    ".cs": "code",
    ".php": "code",
    ".swift": "code",
    ".kt": "code",
    ".scala": "code",
    ".sh": "code",
    ".bash": "code",
    ".zsh": "code",
    ".sql": "code",
    ".json": "code",
    ".yaml": "code",
    ".yml": "code",
    ".toml": "code",
    ".xml": "code",
    ".css": "code",
    ".scss": "code",
    ".rst": "text",
    ".csv": "text",
    ".log": "text",
}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class Document:
    """An ingested document."""

    id: str = ""
    path: str = ""
    title: str = ""
    content: str = ""
    doc_type: str = "text"  # text, markdown, html, pdf, docx, code
    metadata: JsonDict = field(default_factory=dict)
    chunk_count: int = 0
    ingested_at: str = ""
    checksum: str = ""

    def __post_init__(self) -> None:
        if not self.id:
            raw = f"{self.path}{self.title}{time.time_ns()}"
            self.id = hashlib.sha256(raw.encode()).hexdigest()[:16]
        if not self.ingested_at:
            self.ingested_at = datetime.now(timezone.utc).isoformat()
        if self.content and not self.checksum:
            self.checksum = hashlib.sha256(self.content.encode()).hexdigest()[:16]

    def to_dict(self) -> JsonDict:
        return {
            "id": self.id,
            "path": self.path,
            "title": self.title,
            "content": self.content,
            "doc_type": self.doc_type,
            "metadata": self.metadata,
            "chunk_count": self.chunk_count,
            "ingested_at": self.ingested_at,
            "checksum": self.checksum,
        }

    @classmethod
    def from_dict(cls, data: JsonDict) -> "Document":
        return cls(
            id=data.get("id", ""),
            path=data.get("path", ""),
            title=data.get("title", ""),
            content=data.get("content", ""),
            doc_type=data.get("doc_type", "text"),
            metadata=data.get("metadata", {}),
            chunk_count=data.get("chunk_count", 0),
            ingested_at=data.get("ingested_at", ""),
            checksum=data.get("checksum", ""),
        )


@dataclass
class TextChunk:
    """A chunk of text from a document."""

    id: str = ""
    document_id: str = ""
    content: str = ""
    index: int = 0  # chunk index within document
    start_char: int = 0
    end_char: int = 0
    embedding: Optional[EmbeddingVector] = None
    metadata: JsonDict = field(default_factory=dict)

    def __post_init__(self) -> None:
        if not self.id:
            raw = f"{self.document_id}:{self.index}:{self.start_char}"
            self.id = hashlib.sha256(raw.encode()).hexdigest()[:16]


@dataclass
class SearchResult:
    """A single search result with source citation."""

    chunk: TextChunk
    score: float
    source_path: str = ""
    source_title: str = ""
    chunk_index: int = 0
    match_type: str = "hybrid"  # vector | keyword | hybrid

    @property
    def citation(self) -> str:
        """Human-readable source citation."""
        parts: List[str] = []
        if self.source_title:
            parts.append(self.source_title)
        if self.source_path:
            parts.append(f"({self.source_path})")
        parts.append(f"chunk {self.chunk_index}")
        return " — ".join(parts)


@dataclass
class RAGResponse:
    """A RAG query response with context and citations."""

    query: str
    context: str  # formatted context for LLM injection
    results: List[SearchResult]
    citations: List[str]  # unique source citations
    total_chunks_searched: int = 0
    search_time_ms: float = 0.0


# ---------------------------------------------------------------------------
# Embedding backends (shared with store.py, re-implemented here for independence)
# ---------------------------------------------------------------------------


class EmbeddingBackend(Protocol):
    def encode(self, texts: List[str]) -> List[EmbeddingVector]: ...
    @property
    def dimension(self) -> int: ...
    @property
    def name(self) -> str: ...


class SentenceTransformerBackend:
    def __init__(self, model_name: str = DEFAULT_EMBEDDING_MODEL) -> None:
        if not HAS_SENTENCE_TRANSFORMERS:
            raise RuntimeError("sentence-transformers not installed")
        logger.info("RAG: Loading embedding model: %s", model_name)
        self._model = SentenceTransformer(model_name)
        self._dim: int = self._model.get_sentence_embedding_dimension()

    def encode(self, texts: List[str]) -> List[EmbeddingVector]:
        vectors = self._model.encode(texts, show_progress_bar=False, normalize_embeddings=True)
        return [v.tolist() for v in vectors]

    @property
    def dimension(self) -> int:
        return self._dim

    @property
    def name(self) -> str:
        return f"sentence-transformers({DEFAULT_EMBEDDING_MODEL})"


class TFIDFBackend:
    def __init__(self, dim: int = 384) -> None:
        self._dim = dim
        self._idf: Dict[str, float] = {}
        self._vocab: Dict[str, int] = {}
        self._doc_count = 0

    def _tokenize(self, text: str) -> List[str]:
        return text.lower().split()

    def _build_vocab(self, texts: List[str]) -> None:
        df: Counter = Counter()
        for text in texts:
            for t in set(self._tokenize(text)):
                df[t] += 1
        self._doc_count = len(texts)
        for term, freq in df.items():
            if term not in self._vocab:
                self._vocab[term] = len(self._vocab) % self._dim
            self._idf[term] = math.log((self._doc_count + 1) / (freq + 1)) + 1

    def _vectorize(self, text: str) -> EmbeddingVector:
        vec = [0.0] * self._dim
        tokens = self._tokenize(text)
        tf: Counter = Counter(tokens)
        total = len(tokens) or 1
        for term, count in tf.items():
            idx = self._vocab.get(term)
            if idx is not None:
                vec[idx] = (count / total) * self._idf.get(term, 1.0)
        norm = math.sqrt(sum(x * x for x in vec)) or 1.0
        return [x / norm for x in vec]

    def encode(self, texts: List[str]) -> List[EmbeddingVector]:
        self._build_vocab(texts)
        return [self._vectorize(t) for t in texts]

    @property
    def dimension(self) -> int:
        return self._dim

    @property
    def name(self) -> str:
        return "tfidf-fallback"


def _get_embedding_backend(model_name: str = DEFAULT_EMBEDDING_MODEL) -> EmbeddingBackend:
    if HAS_SENTENCE_TRANSFORMERS:
        try:
            return SentenceTransformerBackend(model_name)
        except Exception as exc:
            logger.warning("sentence-transformers failed (%s), using TF-IDF", exc)
    return TFIDFBackend()


# ---------------------------------------------------------------------------
# Vector store backends (shared pattern with store.py)
# ---------------------------------------------------------------------------


class VectorStore(Protocol):
    def add(self, ids: List[str], vectors: List[EmbeddingVector], metadatas: List[JsonDict]) -> None: ...
    def query(self, vector: EmbeddingVector, top_k: int) -> Tuple[List[str], List[float]]: ...
    def delete(self, ids: List[str]) -> None: ...
    def count(self) -> int: ...
    def save(self) -> None: ...


class InMemoryVectorStore:
    def __init__(self) -> None:
        self._ids: List[str] = []
        self._vectors: List[EmbeddingVector] = []
        self._metadatas: List[JsonDict] = []

    def add(self, ids: List[str], vectors: List[EmbeddingVector], metadatas: List[JsonDict]) -> None:
        for id_, vec, meta in zip(ids, vectors, metadatas):
            self._ids.append(id_)
            self._vectors.append(vec)
            self._metadatas.append(meta)

    def query(self, vector: EmbeddingVector, top_k: int = DEFAULT_TOP_K) -> Tuple[List[str], List[float]]:
        if not self._vectors:
            return [], []
        if HAS_NUMPY:
            q = np.array(vector, dtype=np.float32)
            mat = np.array(self._vectors, dtype=np.float32)
            sims = mat @ q
            k = min(top_k, len(self._ids))
            idxs = np.argpartition(-sims, k)[:k]
            idxs = idxs[np.argsort(-sims[idxs])]
            return [self._ids[i] for i in idxs], [float(sims[i]) for i in idxs]
        scores = [(sum(a * b for a, b in zip(vector, v)), i) for i, v in enumerate(self._vectors)]
        scores.sort(reverse=True)
        top = scores[:top_k]
        return [self._ids[i] for _, i in top], [s for s, _ in top]

    def delete(self, ids: List[str]) -> None:
        id_set = set(ids)
        keep = [(i, v, m) for i, v, m in zip(self._ids, self._vectors, self._metadatas) if i not in id_set]
        if keep:
            self._ids, self._vectors, self._metadatas = zip(*keep)
            self._ids, self._vectors, self._metadatas = list(self._ids), list(self._vectors), list(self._metadatas)
        else:
            self._ids, self._vectors, self._metadatas = [], [], []

    def count(self) -> int:
        return len(self._ids)

    def save(self) -> None:
        pass


class ChromaVectorStore:
    COLLECTION_NAME = "rally_rag_documents"

    def __init__(self, persist_dir: str) -> None:
        if not HAS_CHROMADB:
            raise RuntimeError("chromadb not installed")
        os.makedirs(persist_dir, exist_ok=True)
        self._client = chromadb.Client(
            ChromaSettings(
                chroma_db_impl="duckdb+parquet",
                persist_directory=persist_dir,
                anonymized_telemetry=False,
            )
        )
        self._collection = self._client.get_or_create_collection(
            name=self.COLLECTION_NAME,
            metadata={"hnsw:space": "cosine"},
        )

    def add(self, ids: List[str], vectors: List[EmbeddingVector], metadatas: List[JsonDict]) -> None:
        self._collection.upsert(ids=ids, embeddings=vectors, metadatas=metadatas)

    def query(self, vector: EmbeddingVector, top_k: int = DEFAULT_TOP_K) -> Tuple[List[str], List[float]]:
        results = self._collection.query(query_embeddings=[vector], n_results=top_k)
        ids = results.get("ids", [[]])[0]
        dists = results.get("distances", [[]])[0]
        sims = [1.0 - d for d in dists]
        return ids, sims

    def delete(self, ids: List[str]) -> None:
        if ids:
            self._collection.delete(ids=ids)

    def count(self) -> int:
        return self._collection.count()

    def save(self) -> None:
        self._client.persist()


def _get_vector_store(persist_dir: str) -> VectorStore:
    if HAS_CHROMADB:
        try:
            return ChromaVectorStore(persist_dir)
        except Exception as exc:
            logger.warning("ChromaDB init failed (%s), using in-memory", exc)
    return InMemoryVectorStore()


# ---------------------------------------------------------------------------
# BM25 keyword scorer
# ---------------------------------------------------------------------------


class BM25Scorer:
    def __init__(self, k1: float = 1.5, b: float = 0.75) -> None:
        self._k1 = k1
        self._b = b
        self._doc_freqs: Dict[str, int] = {}
        self._doc_lens: List[int] = []
        self._avg_dl: float = 0.0
        self._num_docs: int = 0
        self._tf: List[Dict[str, int]] = []

    def index(self, documents: List[str]) -> None:
        self._num_docs = len(documents)
        self._doc_freqs.clear()
        self._doc_lens.clear()
        self._tf.clear()
        total_len = 0
        for doc in documents:
            tokens = doc.lower().split()
            self._doc_lens.append(len(tokens))
            total_len += len(tokens)
            tf: Counter = Counter(tokens)
            self._tf.append(dict(tf))
            for term in set(tokens):
                self._doc_freqs[term] = self._doc_freqs.get(term, 0) + 1
        self._avg_dl = total_len / self._num_docs if self._num_docs else 1.0

    def score(self, query: str, doc_idx: int) -> float:
        tokens = query.lower().split()
        dl = self._doc_lens[doc_idx]
        tf_map = self._tf[doc_idx]
        score = 0.0
        for term in tokens:
            if term not in tf_map:
                continue
            tf = tf_map[term]
            df = self._doc_freqs.get(term, 0)
            idf = math.log((self._num_docs - df + 0.5) / (df + 0.5) + 1.0)
            num = tf * (self._k1 + 1)
            denom = tf + self._k1 * (1 - self._b + self._b * dl / self._avg_dl)
            score += idf * (num / denom)
        return score

    def score_all(self, query: str) -> List[Tuple[int, float]]:
        """Score all indexed documents. Returns list of (index, score)."""
        results = []
        for i in range(self._num_docs):
            s = self.score(query, i)
            if s > 0:
                results.append((i, s))
        results.sort(key=lambda x: x[1], reverse=True)
        return results


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------


def _extract_text_from_file(path: str, doc_type: str) -> str:
    """Extract plain text from various file formats."""
    try:
        if doc_type == "pdf":
            return _extract_pdf(path)
        elif doc_type == "docx":
            return _extract_docx(path)
        elif doc_type == "html":
            return _extract_html(path)
        else:
            # Plain text / markdown / code
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
    except Exception as exc:
        logger.error("Failed to extract text from %s: %s", path, exc)
        return ""


def _extract_pdf(path: str) -> str:
    """Extract text from PDF (requires PyPDF2 or pdfplumber, with fallback)."""
    try:
        import PyPDF2

        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
    except ImportError:
        pass

    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            return "\n\n".join(pages)
    except ImportError:
        pass

    logger.warning("No PDF library available (install PyPDF2 or pdfplumber)")
    return ""


def _extract_docx(path: str) -> str:
    """Extract text from DOCX (requires python-docx)."""
    try:
        import docx

        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed; cannot read .docx")
        return ""


def _extract_html(path: str) -> str:
    """Extract text from HTML (uses html.parser, no external deps)."""
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._parts: List[str] = []
            self._skip = False

        def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
            if tag in ("script", "style", "noscript"):
                self._skip = True

        def handle_endtag(self, tag: str) -> None:
            if tag in ("script", "style", "noscript"):
                self._skip = False
            if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"):
                self._parts.append("\n")

        def handle_data(self, data: str) -> None:
            if not self._skip:
                self._parts.append(data)

    with open(path, "r", encoding="utf-8", errors="replace") as f:
        html = f.read()

    extractor = _TextExtractor()
    extractor.feed(html)
    text = "".join(extractor._parts)
    # Collapse whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _detect_doc_type(path: str) -> str:
    """Detect document type from file extension."""
    ext = Path(path).suffix.lower()
    category = SUPPORTED_EXTENSIONS.get(ext)
    if category:
        return category
    # Try to read as text
    return "text"


def _title_from_path(path: str) -> str:
    """Derive a title from file path."""
    p = Path(path)
    return p.stem.replace("_", " ").replace("-", " ").title()


# ---------------------------------------------------------------------------
# Text chunking
# ---------------------------------------------------------------------------


def chunk_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> List[Tuple[str, int, int]]:
    """Split text into overlapping chunks.

    Returns list of (chunk_text, start_char, end_char).
    Tries to break at paragraph or sentence boundaries when possible.
    """
    if not text.strip():
        return []

    if len(text) <= chunk_size:
        return [(text, 0, len(text))]

    chunks: List[Tuple[str, int, int]] = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        # Try to break at a paragraph boundary
        if end < len(text):
            # Look for paragraph break near the end
            para_break = text.rfind("\n\n", start, end)
            if para_break > start + chunk_size // 2:
                end = para_break + 2
            else:
                # Try sentence boundary
                for delim in (". ", "! ", "? ", "\n"):
                    sent_break = text.rfind(delim, start + chunk_size // 2, end)
                    if sent_break > start:
                        end = sent_break + len(delim)
                        break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append((chunk, start, end))

        # Advance with overlap
        start = max(start + 1, end - chunk_overlap)

    return chunks


# ---------------------------------------------------------------------------
# RAG Pipeline
# ---------------------------------------------------------------------------


class RAGPipeline:
    """Full RAG pipeline: ingest → chunk → embed → store → retrieve → cite.

    Usage::

        rag = RAGPipeline(data_dir="/path/to/index")
        await rag.ingest_file("doc.pdf")
        await rag.ingest_directory("./docs/")
        response = await rag.query("What is the main topic?")
        print(response.context)   # inject into LLM
        print(response.citations) # source references
    """

    def __init__(
        self,
        *,
        data_dir: Optional[str] = None,
        embedding_model: str = DEFAULT_EMBEDDING_MODEL,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        top_k: int = DEFAULT_TOP_K,
    ) -> None:
        self._data_dir = data_dir or os.path.expanduser("~/.rally-agent/rag")
        os.makedirs(self._data_dir, exist_ok=True)

        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._top_k = top_k

        # Internal stores
        self._documents: Dict[str, Document] = {}
        self._chunks: Dict[str, TextChunk] = {}
        self._chunk_order: List[str] = []  # for BM25 index alignment

        # Embedding + vector
        self._embedder: Optional[EmbeddingBackend] = None
        self._embedding_model_name = embedding_model
        self._vector_store: Optional[VectorStore] = None

        # BM25
        self._bm25 = BM25Scorer()
        self._bm25_dirty = True

        # Persist paths
        self._index_file = os.path.join(self._data_dir, "rag_index.json")
        self._embed_file = os.path.join(self._data_dir, f"rag{EMBEDDING_INDEX_EXT}")

        # Load persisted index
        self._load()

        logger.info(
            "RAGPipeline initialised: %d documents, %d chunks",
            len(self._documents),
            len(self._chunks),
        )

    # -- lazy init ------------------------------------------------------------

    def _ensure_embedder(self) -> EmbeddingBackend:
        if self._embedder is None:
            self._embedder = _get_embedding_backend(self._embedding_model_name)
        return self._embedder

    def _ensure_vector_store(self) -> VectorStore:
        if self._vector_store is None:
            persist_dir = os.path.join(self._data_dir, "vectorstore")
            self._vector_store = _get_vector_store(persist_dir)
            self._reindex_vectors()
        return self._vector_store

    def _reindex_vectors(self) -> None:
        ids, vecs, metas = [], [], []
        for chunk in self._chunks.values():
            if chunk.embedding is not None:
                ids.append(chunk.id)
                vecs.append(chunk.embedding)
                metas.append({
                    "document_id": chunk.document_id,
                    "chunk_index": chunk.index,
                })
        if ids:
            self._vector_store.add(ids, vecs, metas)

    # -- persistence ----------------------------------------------------------

    def _load(self) -> None:
        if not os.path.exists(self._index_file):
            return
        try:
            with open(self._index_file, "r", encoding="utf-8") as f:
                data = json.load(f)
        except (json.JSONDecodeError, OSError) as exc:
            logger.error("Failed to load RAG index: %s", exc)
            return

        for doc_data in data.get("documents", []):
            doc = Document.from_dict(doc_data)
            self._documents[doc.id] = doc

        embeddings = self._load_embeddings()

        for chunk_data in data.get("chunks", []):
            chunk = TextChunk(
                id=chunk_data["id"],
                document_id=chunk_data["document_id"],
                content=chunk_data["content"],
                index=chunk_data.get("index", 0),
                start_char=chunk_data.get("start_char", 0),
                end_char=chunk_data.get("end_char", 0),
                metadata=chunk_data.get("metadata", {}),
            )
            if chunk.id in embeddings:
                chunk.embedding = embeddings[chunk.id]
            self._chunks[chunk.id] = chunk
            self._chunk_order.append(chunk.id)

        self._bm25_dirty = True

    def _load_embeddings(self) -> Dict[str, EmbeddingVector]:
        embeddings: Dict[str, EmbeddingVector] = {}
        if not os.path.exists(self._embed_file):
            return embeddings
        try:
            with open(self._embed_file, "rb") as f:
                count = struct.unpack("<I", f.read(4))[0]
                for _ in range(count):
                    id_len = struct.unpack("<H", f.read(2))[0]
                    cid = f.read(id_len).decode("utf-8")
                    vec_len = struct.unpack("<I", f.read(4))[0]
                    vec = list(struct.unpack(f"<{vec_len}f", f.read(vec_len * 4)))
                    embeddings[cid] = vec
        except Exception as exc:
            logger.warning("Failed to load RAG embeddings: %s", exc)
        return embeddings

    def save(self) -> None:
        """Persist the RAG index to disk."""
        data = {
            "documents": [d.to_dict() for d in self._documents.values()],
            "chunks": [
                {
                    "id": c.id,
                    "document_id": c.document_id,
                    "content": c.content,
                    "index": c.index,
                    "start_char": c.start_char,
                    "end_char": c.end_char,
                    "metadata": c.metadata,
                }
                for c in self._chunks.values()
            ],
        }
        tmp = self._index_file + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        os.replace(tmp, self._index_file)

        self._save_embeddings()

        if self._vector_store:
            self._vector_store.save()

    def _save_embeddings(self) -> None:
        items = [(c.id, c.embedding) for c in self._chunks.values() if c.embedding is not None]
        if not items:
            return
        tmp = self._embed_file + ".tmp"
        with open(tmp, "wb") as f:
            f.write(struct.pack("<I", len(items)))
            for cid, vec in items:
                id_bytes = cid.encode("utf-8")
                f.write(struct.pack("<H", len(id_bytes)))
                f.write(id_bytes)
                f.write(struct.pack("<I", len(vec)))
                f.write(struct.pack(f"<{len(vec)}f", *vec))
        os.replace(tmp, self._embed_file)

    # -- BM25 -----------------------------------------------------------------

    def _ensure_bm25(self) -> None:
        if self._bm25_dirty:
            texts = [self._chunks[cid].content for cid in self._chunk_order if cid in self._chunks]
            self._bm25.index(texts)
            self._bm25_dirty = False

    # -- document ingestion ---------------------------------------------------

    async def ingest_file(
        self,
        path: str,
        *,
        title: Optional[str] = None,
        metadata: Optional[JsonDict] = None,
    ) -> Document:
        """Ingest a single file into the RAG index.

        Supports: PDF, DOCX, HTML, Markdown, plain text, code files.
        Returns the created Document object.
        """
        path = os.path.abspath(path)
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")

        doc_type = _detect_doc_type(path)
        content = _extract_text_from_file(path, doc_type)
        if not content.strip():
            raise ValueError(f"No text extracted from {path}")

        return await self.ingest_text(
            content=content,
            title=title or _title_from_path(path),
            path=path,
            doc_type=doc_type,
            metadata=metadata,
        )

    async def ingest_text(
        self,
        content: str,
        *,
        title: str = "",
        path: str = "",
        doc_type: str = "text",
        metadata: Optional[JsonDict] = None,
    ) -> Document:
        """Ingest raw text content into the RAG index."""
        doc = Document(
            path=path,
            title=title,
            content=content,
            doc_type=doc_type,
            metadata=metadata or {},
        )

        # Check for duplicate by checksum
        for existing in self._documents.values():
            if existing.checksum == doc.checksum:
                logger.info("Document already indexed (checksum match): %s", existing.id)
                return existing

        # Chunk the content
        raw_chunks = chunk_text(content, self._chunk_size, self._chunk_overlap)
        doc.chunk_count = len(raw_chunks)

        # Create chunk objects
        chunks: List[TextChunk] = []
        for i, (text, start, end) in enumerate(raw_chunks):
            chunk = TextChunk(
                document_id=doc.id,
                content=text,
                index=i,
                start_char=start,
                end_char=end,
                metadata={"title": title, "path": path, "doc_type": doc_type},
            )
            chunks.append(chunk)
            self._chunks[chunk.id] = chunk
            self._chunk_order.append(chunk.id)

        # Generate embeddings
        await self._embed_chunks(chunks)

        # Store document
        self._documents[doc.id] = doc
        self._bm25_dirty = True

        # Persist
        self.save()

        logger.info("Ingested document '%s': %d chunks", title or path, len(chunks))
        return doc

    async def ingest_directory(
        self,
        dir_path: str,
        *,
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """Ingest all supported files in a directory."""
        dir_path = os.path.abspath(dir_path)
        if not os.path.isdir(dir_path):
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        allowed = set(extensions) if extensions else set(SUPPORTED_EXTENSIONS.keys())
        documents: List[Document] = []

        pattern = "**/*" if recursive else "*"
        for fpath in sorted(Path(dir_path).glob(pattern)):
            if not fpath.is_file():
                continue
            if fpath.suffix.lower() not in allowed:
                continue
            try:
                doc = await self.ingest_file(str(fpath))
                documents.append(doc)
            except Exception as exc:
                logger.warning("Skipping %s: %s", fpath, exc)

        return documents

    async def _embed_chunks(self, chunks: List[TextChunk]) -> None:
        """Generate and store embeddings for a batch of chunks."""
        embedder = self._ensure_embedder()
        store = self._ensure_vector_store()

        texts = [c.content for c in chunks]
        try:
            vectors = embedder.encode(texts)
        except Exception as exc:
            logger.error("Embedding failed: %s", exc)
            return

        ids, metas = [], []
        for chunk, vec in zip(chunks, vectors):
            chunk.embedding = vec
            ids.append(chunk.id)
            metas.append({
                "document_id": chunk.document_id,
                "chunk_index": chunk.index,
            })

        store.add(ids, vectors, metas)

    # -- document management --------------------------------------------------

    async def remove_document(self, document_id: str) -> bool:
        """Remove a document and all its chunks from the index."""
        doc = self._documents.get(document_id)
        if doc is None:
            return False

        # Find and remove chunks
        chunk_ids = [cid for cid, c in self._chunks.items() if c.document_id == document_id]
        for cid in chunk_ids:
            del self._chunks[cid]
        self._chunk_order = [cid for cid in self._chunk_order if cid not in chunk_ids]

        # Remove from vector store
        if self._vector_store and chunk_ids:
            self._vector_store.delete(chunk_ids)

        del self._documents[document_id]
        self._bm25_dirty = True
        self.save()
        logger.info("Removed document '%s' (%d chunks)", doc.title, len(chunk_ids))
        return True

    async def update_document(self, path: str, **kwargs: Any) -> Document:
        """Re-ingest a document (remove old, add new)."""
        path = os.path.abspath(path)
        # Remove existing
        for doc in list(self._documents.values()):
            if doc.path == path:
                await self.remove_document(doc.id)
                break
        return await self.ingest_file(path, **kwargs)

    def list_documents(self) -> List[Document]:
        """List all indexed documents."""
        return list(self._documents.values())

    # -- query / search -------------------------------------------------------

    async def query(
        self,
        query: str,
        *,
        top_k: Optional[int] = None,
        mode: Literal["hybrid", "vector", "keyword"] = "hybrid",
        include_content: bool = True,
        max_context_chars: int = 8000,
    ) -> RAGResponse:
        """Query the RAG index and return context + citations.

        The ``context`` field in the response is ready to inject into an LLM prompt.
        """
        t0 = time.monotonic()
        k = top_k or self._top_k

        results = await self.search(query, top_k=k * 2, mode=mode)

        # Deduplicate by document (keep best chunk per doc for diversity)
        seen_docs: set = set()
        deduped: List[SearchResult] = []
        for r in results:
            if r.chunk.document_id not in seen_docs or len(deduped) < k:
                deduped.append(r)
                seen_docs.add(r.chunk.document_id)
            if len(deduped) >= k:
                break

        # Build context
        context_parts: List[str] = []
        used_chars = 0
        for r in deduped:
            if not include_content:
                continue
            block = f"[{r.citation}]\n{r.chunk.content}"
            if used_chars + len(block) > max_context_chars:
                break
            context_parts.append(block)
            used_chars += len(block) + 2

        context = "\n\n---\n\n".join(context_parts) if context_parts else ""

        # Citations
        citations = list(dict.fromkeys(r.citation for r in deduped))

        elapsed = (time.monotonic() - t0) * 1000

        return RAGResponse(
            query=query,
            context=context,
            results=deduped,
            citations=citations,
            total_chunks_searched=len(self._chunks),
            search_time_ms=round(elapsed, 2),
        )

    async def search(
        self,
        query: str,
        *,
        top_k: int = DEFAULT_TOP_K,
        mode: Literal["hybrid", "vector", "keyword"] = "hybrid",
    ) -> List[SearchResult]:
        """Raw search returning ranked SearchResult objects."""
        results: Dict[str, SearchResult] = {}
        chunk_map = {cid: self._chunks[cid] for cid in self._chunk_order if cid in self._chunks}

        # Vector search
        if mode in ("hybrid", "vector"):
            try:
                vec_results = await self._vector_search(query, top_k * 3)
                for chunk_id, sim in vec_results:
                    chunk = chunk_map.get(chunk_id)
                    if chunk is None:
                        continue
                    doc = self._documents.get(chunk.document_id)
                    results[chunk_id] = SearchResult(
                        chunk=chunk,
                        score=sim,
                        source_path=doc.path if doc else "",
                        source_title=doc.title if doc else "",
                        chunk_index=chunk.index,
                        match_type="vector",
                    )
            except Exception as exc:
                logger.debug("Vector search unavailable: %s", exc)

        # BM25 keyword search
        if mode in ("hybrid", "keyword"):
            self._ensure_bm25()
            bm25_results = self._bm25.score_all(query)
            if bm25_results:
                max_bm25 = bm25_results[0][1]
                # Map BM25 index back to chunk ids
                chunk_id_list = [cid for cid in self._chunk_order if cid in self._chunks]
                for idx, s in bm25_results[: top_k * 3]:
                    if idx >= len(chunk_id_list):
                        continue
                    cid = chunk_id_list[idx]
                    normed = s / max_bm25 if max_bm25 > 0 else 0
                    chunk = chunk_map.get(cid)
                    if chunk is None:
                        continue
                    if cid in results:
                        old = results[cid]
                        old.score = old.score * 0.6 + normed * 0.4
                        old.match_type = "hybrid"
                    else:
                        doc = self._documents.get(chunk.document_id)
                        results[cid] = SearchResult(
                            chunk=chunk,
                            score=normed * 0.4,
                            source_path=doc.path if doc else "",
                            source_title=doc.title if doc else "",
                            chunk_index=chunk.index,
                            match_type="keyword",
                        )

        sorted_results = sorted(results.values(), key=lambda r: r.score, reverse=True)[:top_k]
        return sorted_results

    async def _vector_search(self, query: str, limit: int) -> List[Tuple[str, float]]:
        store = self._ensure_vector_store()
        embedder = self._ensure_embedder()

        # Embed pending chunks
        pending = [c for c in self._chunks.values() if c.embedding is None]
        if pending:
            await self._embed_chunks(pending)

        if store.count() == 0:
            return []

        q_vec = embedder.encode([query])[0]
        ids, scores = store.query(q_vec, limit)
        return list(zip(ids, scores))

    # -- stats ----------------------------------------------------------------

    def stats(self) -> Dict[str, Any]:
        """Return pipeline statistics."""
        total_chunks = len(self._chunks)
        embedded = sum(1 for c in self._chunks.values() if c.embedding is not None)
        return {
            "documents": len(self._documents),
            "total_chunks": total_chunks,
            "embedded_chunks": embedded,
            "vector_store_count": self._vector_store.count() if self._vector_store else 0,
            "embedding_backend": self._ensure_embedder().name,
            "chunk_size": self._chunk_size,
            "chunk_overlap": self._chunk_overlap,
            "top_k": self._top_k,
        }

    def show_stats(self) -> None:
        """Pretty-print pipeline stats (CLI)."""
        try:
            from cli.theme import Theme, console

            s = self.stats()
            table = Theme.create_table("📚 RAG Pipeline Statistics")
            table.add_column("Metric", style="cyan")
            table.add_column("Value", style="neon_green")
            for k, v in s.items():
                table.add_row(k.replace("_", " ").title(), str(v))
            console.print()
            console.print(table)
            console.print()
        except ImportError:
            s = self.stats()
            print(f"📚 RAG: {s['documents']} docs, {s['total_chunks']} chunks, {s['embedded_chunks']} embedded")
